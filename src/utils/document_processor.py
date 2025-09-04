"""
Document Processing Utilities
Handles extraction and preprocessing of text from various document formats
"""

import os
import re
import logging
from typing import Dict, List, Optional, Tuple
import PyPDF2
from docx import Document
import pandas as pd
from io import StringIO, BytesIO

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    A comprehensive document processor that extracts text from multiple file formats
    and provides preprocessing capabilities.
    """
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.txt', '.csv', '.xlsx'}
    
    def __init__(self):
        self.processed_documents = {}
    
    def extract_text(self, file_path: str, file_content: bytes = None) -> Dict[str, any]:
        """
        Extract text from various document formats
        
        Args:
            file_path: Path to the document or filename with extension
            file_content: Raw file content as bytes (for uploaded files)
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path, file_content)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path, file_content)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path, file_content)
            elif file_extension in ['.csv', '.xlsx']:
                return self._extract_from_spreadsheet(file_path, file_content)
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                'text': '',
                'metadata': {
                    'error': str(e),
                    'file_type': file_extension,
                    'word_count': 0,
                    'char_count': 0
                }
            }
    
    def _extract_from_pdf(self, file_path: str, file_content: bytes = None) -> Dict:
        """Extract text from PDF files"""
        text = ""
        page_count = 0
        
        try:
            if file_content:
                pdf_file = BytesIO(file_content)
            else:
                pdf_file = open(file_path, 'rb')
            
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            page_count = len(pdf_reader.pages)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            if not file_content:
                pdf_file.close()
                
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            text = ""
        
        return {
            'text': text.strip(),
            'metadata': {
                'file_type': 'pdf',
                'pages': page_count,
                'word_count': len(text.split()),
                'char_count': len(text)
            }
        }
    
    def _extract_from_docx(self, file_path: str, file_content: bytes = None) -> Dict:
        """Extract text from DOCX files"""
        text = ""
        paragraph_count = 0
        
        try:
            if file_content:
                doc = Document(BytesIO(file_content))
            else:
                doc = Document(file_path)
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                paragraph_count += 1
                
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            text = ""
        
        return {
            'text': text.strip(),
            'metadata': {
                'file_type': 'docx',
                'paragraphs': paragraph_count,
                'word_count': len(text.split()),
                'char_count': len(text)
            }
        }
    
    def _extract_from_txt(self, file_path: str, file_content: bytes = None) -> Dict:
        """Extract text from TXT files"""
        text = ""
        
        try:
            if file_content:
                text = file_content.decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
                    
        except UnicodeDecodeError:
            try:
                if file_content:
                    text = file_content.decode('latin-1')
                else:
                    with open(file_path, 'r', encoding='latin-1') as file:
                        text = file.read()
            except Exception as e:
                logger.error(f"TXT extraction error: {str(e)}")
                text = ""
        
        return {
            'text': text.strip(),
            'metadata': {
                'file_type': 'txt',
                'lines': len(text.split('\n')),
                'word_count': len(text.split()),
                'char_count': len(text)
            }
        }
    
    def _extract_from_spreadsheet(self, file_path: str, file_content: bytes = None) -> Dict:
        """Extract text from CSV/Excel files"""
        text = ""
        row_count = 0
        col_count = 0
        
        try:
            if file_path.endswith('.csv'):
                if file_content:
                    df = pd.read_csv(StringIO(file_content.decode('utf-8')))
                else:
                    df = pd.read_csv(file_path)
            else:  # Excel
                if file_content:
                    df = pd.read_excel(BytesIO(file_content))
                else:
                    df = pd.read_excel(file_path)
            
            # Convert dataframe to text representation
            text = df.to_string(index=False)
            row_count, col_count = df.shape
            
        except Exception as e:
            logger.error(f"Spreadsheet extraction error: {str(e)}")
            text = ""
        
        return {
            'text': text.strip(),
            'metadata': {
                'file_type': 'spreadsheet',
                'rows': row_count,
                'columns': col_count,
                'word_count': len(text.split()),
                'char_count': len(text)
            }
        }
    
    def preprocess_text(self, text: str, options: Dict = None) -> str:
        """
        Preprocess text for analysis
        
        Args:
            text: Raw text to preprocess
            options: Dictionary of preprocessing options
            
        Returns:
            Preprocessed text
        """
        if not options:
            options = {
                'remove_extra_whitespace': True,
                'remove_special_chars': False,
                'lowercase': False,
                'remove_numbers': False
            }
        
        processed_text = text
        
        # Remove extra whitespace
        if options.get('remove_extra_whitespace', True):
            processed_text = re.sub(r'\s+', ' ', processed_text).strip()
        
        # Remove special characters (keep basic punctuation)
        if options.get('remove_special_chars', False):
            processed_text = re.sub(r'[^a-zA-Z0-9\s\.\,\!\?\;\:]', '', processed_text)
        
        # Convert to lowercase
        if options.get('lowercase', False):
            processed_text = processed_text.lower()
        
        # Remove numbers
        if options.get('remove_numbers', False):
            processed_text = re.sub(r'\d+', '', processed_text)
        
        return processed_text
    
    def get_document_stats(self, text: str) -> Dict:
        """Get comprehensive statistics about the document"""
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        paragraphs = text.split('\n\n')
        
        # Calculate average word length
        avg_word_length = sum(len(word) for word in words) / len(words) if words else 0
        
        # Calculate readability (simplified)
        avg_sentence_length = len(words) / len([s for s in sentences if s.strip()]) if sentences else 0
        
        return {
            'word_count': len(words),
            'sentence_count': len([s for s in sentences if s.strip()]),
            'paragraph_count': len([p for p in paragraphs if p.strip()]),
            'character_count': len(text),
            'avg_word_length': round(avg_word_length, 2),
            'avg_sentence_length': round(avg_sentence_length, 2),
            'unique_words': len(set(word.lower() for word in words if word.isalpha()))
        }
